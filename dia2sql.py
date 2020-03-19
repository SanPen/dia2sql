import os
import gzip
from xml.etree import cElementTree as ElementTree
import pandas as pd
from docx import Document  # from python-docx package


class XmlDictConfig(dict):
    """
    Note: need to add a root into if no exising
    Example usage:
    >>> tree = ElementTree.parse('your_file.xml')
    >>> root = tree.getroot()
    >>> xmldict = XmlDictConfig(root)
    Or, if you want to use an XML string:
    >>> root = ElementTree.XML(xml_string)
    >>> xmldict = XmlDictConfig(root)
    And then use xmldict for what it is... a dict.
    """
    def __init__(self, parent_element, text_to_remove=''):
        """

        :param parent_element:
        :param text_to_remove:
        """
        self.text_to_remove = text_to_remove

        if parent_element.items():
            self.update_shim(dict(parent_element.items()))

        for element in parent_element:

            tag = element.tag.replace(self.text_to_remove, '')

            if len(element):

                a_dict = XmlDictConfig(element, self.text_to_remove)

                self.update_shim({tag: a_dict})

            elif element.items():    # items() is specialy for attribtes

                element_attrib = element.items()

                if element.text:
                    element_attrib.append((tag, element.text))     # add tag:text if there exist

                self.update_shim({tag: dict(element_attrib)})
            else:
                self.update_shim({tag: element.text})

    def update_shim(self, a_dict):
        """

        :param a_dict:
        :return:
        """
        for og_key in a_dict.keys():   # keys() includes tag and attributes

            key = og_key.replace(self.text_to_remove, '')

            if key in self:
                value = self.pop(key)
                if type(value) is not list:
                    list_of_dicts = list()
                    list_of_dicts.append(value)
                    list_of_dicts.append(a_dict[key])
                    self.update({key: list_of_dicts})
                else:
                    value.append(a_dict[key])
                    self.update({key: value})
            else:
                self.update({key: a_dict[key]})  # it was self.update(aDict)


def fix_name(txt, is_name=True):

    chars = {'á': 'a',
             'é': 'e',
             'í': 'i',
             'ó': 'o',
             'ú': 'u',
             '?': '',
             '¿': '',
             'ñ': 'n'}

    if is_name:
        txt2 = txt.strip().replace(' ', '_').lower()
    else:
        txt2 = txt

    for search, replace in chars.items():
        txt2 = txt2.replace(search, replace)
    return txt2


def parse_xml_file(fname):
    """

    :param fname:
    :return:
    """
    f = gzip.open(fname, 'rb')
    content = f.read()

    # parse string as xml
    root = ElementTree.fromstring(content)

    # pass the XML file to a dictionary
    xml_dict = XmlDictConfig(root, text_to_remove='{http://www.lysator.liu.se/~alla/dia/}')

    return xml_dict


class DBAttribute:

    def __init__(self, lst):
        """
        DB attribute
        :param lst: list of xml attributes to be parsed here
        """
        self.name = ''
        self.comment = ''
        self.type = ''
        self.is_primary_key = False
        self.is_nullable = False
        self.is_unique = False

        boold = {'true': True, 'false': False}

        for elm in lst:
            if 'name' in elm.keys():
                if elm['name'] == 'name':
                    self.name = fix_name(elm['string'].replace('#', '').strip().replace(' ', '_'))
                elif elm['name'] == 'type':
                    self.type = elm['string'].replace('#', '')
                elif elm['name'] == 'comment':
                    self.comment = elm['string'].replace('#', '')
                elif elm['name'] == 'primary_key':
                    self.is_primary_key = boold[elm['boolean']['val']]
                elif elm['name'] == 'nullable':
                    self.is_nullable = boold[elm['boolean']['val']]
                elif elm['name'] == 'unique':
                    self.is_unique = boold[elm['boolean']['val']]

    def fix(self):

        t = self.type.strip().lower()

        if t == 'int':
            self.type = 'INTEGER'

        elif t == 'float':
            self.type = 'REAL'

        elif t == 'double':
            self.type = 'REAL'

    def check(self):
        """
        Check attribute
        :return: string with the warnings
        """
        val = ''
        t = self.type.strip().lower()
        if 'varchar' in t and '(' not in t:
            val += 'The type of ' + self.name + ' requires length i.e VARCHAR(10)'
        elif 'numeric' in t and '(' not in t:
            val += 'The type of ' + self.name + ' requires length i.e NUMERIC(10, 2)'
        elif t == '':
            val += 'The attribute ' + self.name + ' has no type!'

        return val

    def to_sql(self, include_pk=False):
        """

        """
        val = self.name.strip().lower().replace(' ', '_') + ' ' + str(self.type)

        if self.is_primary_key and include_pk:
            val += ' PRIMARY KEY'

        if self.is_nullable:
            val += ' NOT NULL'

        if self.is_unique:
            val += ' UNIQUE'

        if self.comment != "":
            val += '    /* ' + fix_name(self.comment, is_name=False).replace('\n', ' ') + ' */'

        return val

    def __str__(self):
        return self.name


class DiaRelationship:

    def __init__(self, xml_data):
        """

        :param table:
        """
        self.xml = xml_data
        self.id_from = xml_data['connections']['connection'][0]['to']
        self.id_from_connection = xml_data['connections']['connection'][0]['connection']
        self.id_to = xml_data['connections']['connection'][1]['to']
        self.id_to_connection = xml_data['connections']['connection'][1]['connection']
        self.number_from = ''
        self.number_to = ''

        for attr in xml_data['attribute']:
            if attr['name'] == 'start_point_desc':
                self.number_from = attr['string'].replace('#', '')
            elif attr['name'] == 'end_point_desc':
                self.number_to = attr['string'].replace('#', '')

    def __str__(self):
        a = self.id_from + ':' + self.id_from_connection + '(' + self.number_from + ')'
        b = self.id_to + ':' + self.id_to_connection + '(' + self.number_to + ')'
        return a + '->' + b


class DBRelationship:

    def __init__(self, tbl_from, tbl_to, attributes_from, attributes_to):
        """

        :param tbl_from:
        :param tbl_to:
        :param attributes_from:
        :param attributes_to:
        """

        self.tbl_from = tbl_from
        self.tbl_to = tbl_to
        self.attributes_from = attributes_from
        self.attributes_to = attributes_to
        self.name = tbl_from.name + '_' + tbl_to.name

    def __str__(self):
        return self.to_sql()

    def to_sql(self, k=''):
        """
        Generate the SQL equivalent of this constraint
        :param k: some value to attach to the restriction ideally the restriction number
        :return: SQL restriction code
        """
        af = ','.join([e.name for e in self.attributes_from])
        at = ','.join([e.name for e in self.attributes_to])        
        name = 'R' + str(k) + '_' + self.tbl_to.name
        return 'CONSTRAINT ' + name + ' FOREIGN KEY (' + af + ') REFERENCES ' + self.tbl_to.name + '(' + at + ')'


class DBTable:

    def __init__(self, xml_table):
        """
        DB Table
        :param xml_table: xml Table to be parsed here
        """

        self.xml = xml_table

        self.id = xml_table['id']

        self.order = 0

        self.name = ''

        self.comment = ''

        self.attributes = list()

        self.pk = list()

        self.relationships = list()

        # parse attributes:
        for attr in xml_table['attribute']:
            if attr['name'] == 'name':
                self.name = fix_name(attr['string'].replace('#', ''))

            elif attr['name'] == 'comment':
                self.comment = attr['string'].replace('#', '')

            elif attr['name'] == 'attributes':

                if isinstance(attr['composite'], list):
                    for elm in attr['composite']:
                        attribute = DBAttribute(elm['attribute'])
                        self.attributes.append(attribute)

                        if attribute.is_primary_key:
                            self.pk.append(attribute)
                else:
                    attribute = DBAttribute(attr['composite']['attribute'])
                    self.attributes.append(attribute)

                    if attribute.is_primary_key:
                        self.pk.append(attribute)

    def get_parent_tables(self):
        """
        Get list of the tables that this table has foreign keys to
        By doing so, it increments the order of the precedent tables
        :return: list
        """
        tbls = list()
        for rel in self.relationships:
            rel.tbl_to.order += 1
            tbls.append(rel.tbl_to)
        return tbls

    def fix_attribute_errors(self):
        """
        Attempt to fix common mistakes
        """
        for i, attr in enumerate(self.attributes):
            attr.fix()

    def get_errors(self):
        val = ''

        if len(self.pk) == 0:
            val += '\tThere is no primary key!\n'

        # attribute related
        for i, attr in enumerate(self.attributes):
            err = attr.check()
            if err != '':
                val += '\t' + err + '\n'

        val2 = ''

        if val != '':
            val2 = 'Errors:\n' + val

        return val2

    def to_sql(self):
        """
        CREATE TABLE table_name (
            column1 datatype,
            column2 datatype,
            column3 datatype,
           ....
        );
        :return: string
        """
        tab = ' ' * 4

        self.fix_attribute_errors()

        cmnt = fix_name(self.comment, is_name=False)
        val = '/*\nComments:\n' + tab + cmnt.replace('\n', tab + '\n') + '\n' + self.get_errors() + '*/\n'

        val += 'CREATE TABLE ' + self.name + ' (\n'

        # add attributes
        for i, attr in enumerate(self.attributes):
            val += tab + attr.to_sql() + ',\n'

        # add the primary keys
        pk = [a.name for a in self.pk]
        val += tab + 'PRIMARY KEY (' + ', '.join(pk) + ')'

        # add the foreign keys
        if len(self.relationships) > 0:

            val += ',\n'

            for k, rel in enumerate(self.relationships):
                val += tab + rel.to_sql(k)

                if k < len(self.relationships) - 1:
                    val += ',\n'
                else:
                    val += '\n'
        else:
            # remove the last comma
            val += '\n'

        val += ');\n\n'
        return val

    def __str__(self):
        return self.id + ':' + self.name + ':' + str(self.order)


class DBModel:

    def __init__(self, fname=None):

        self.tables = list()

        self.relations = list()
        
        self.sql_code = ''
        
        if os.path.exists(fname):
            self.parse_file(fname)            
            self.output_file = os.path.basename(fname).replace('.dia', '.sql')
        else:
            self.output_file = 'code.sql'
        
    def parse_file(self, fname):
        
        dia_xml = parse_xml_file(fname)

        # get the tables
        for elm in dia_xml['layer']['object']:
            if elm['type'] == 'Database - Table':
                self.tables.append(DBTable(elm))
            elif elm['type'] == 'Database - Reference':
                self.relations.append(DiaRelationship(elm))
    
        self.find_relationships()
        
    def writer(self, txt):
        
        self.sql_code += txt

    def find_relationships(self):
        """

        :return:
        """

        table_dict = {t.id: t for t in self.tables}

        # find relationships
        for rel in self.relations:

            tbl_f = table_dict[rel.id_from]
            tbl_t = table_dict[rel.id_to]

            pk_attributes = [elem.name for elem in tbl_f.pk]
            tbl2_attributes = [e.name for e in tbl_t.attributes]
            # related = all(elem in tbl2_attributes for elem in pk_attributes)
            attributes_to = list()
            for i, elem in enumerate(pk_attributes):
                if elem in tbl2_attributes:
                    j = tbl2_attributes.index(elem)
                    attributes_to.append(tbl_t.attributes[j])

            # create relationship
            r = DBRelationship(tbl_from=tbl_f, 
                               tbl_to=tbl_t, 
                               attributes_from=tbl_f.pk, 
                               attributes_to=attributes_to)
            
            tbl_f.relationships.append(r)

            if len(tbl_f.pk) != len(attributes_to):
                self.writer('/*INVALID RELATIONSHIIP:\n\t' + str(r) + '*/')

        # find the table orders
        for tbl in self.tables:
            tbl.get_parent_tables()

        # sort tables
        self.tables.sort(key=lambda x: x.order, reverse=True)

    def to_sql(self, clear=False):
        """
        Create SQL statement for all the tables
        :return: string
        """
        if clear:
            self.sql_code = ''
        
        self.sql_code += '\n' * 3
        self.sql_code += '/* DIA 2 SQL code generation */'
        self.sql_code += '\n' * 3
        
        # Create code for each table        
        for tbl in self.tables:
            self.sql_code += tbl.to_sql()
            
        return self.sql_code

    def to_ms_word(self, file_name=None):
        """
        Export the Model to a MS Word document
        :param: file_name: Name of the file
        """
        document = Document()

        document.add_heading('Especificación ', 1)

        for tbl in self.tables:

            # one section per table
            document.add_heading(tbl.name, level=2)

            # add the comment
            p = document.add_paragraph(tbl.comment)

            document.add_heading('Atributos', level=3)

            # add the attributes in a table
            doc_table = document.add_table(rows=len(tbl.attributes) + 1, cols=5)

            hdr_cells = doc_table.rows[0].cells
            hdr_cells[0].text = 'Atributo'
            hdr_cells[1].text = 'Tipo'
            hdr_cells[2].text = 'Clave Primaria'
            hdr_cells[3].text = 'Único'
            hdr_cells[4].text = 'Comentario'

            for i, attr in enumerate(tbl.attributes):
                row_cells = doc_table.rows[i+1].cells
                row_cells[0].text = attr.name  # 'Atributo'
                row_cells[1].text = str(attr.type)  # 'Tipo'
                row_cells[2].text = str(attr.is_primary_key)  # 'Clave Primaria'
                row_cells[3].text = str(attr.is_unique)  # 'Único'
                row_cells[4].text = attr.comment  # 'Comentario'

            # add the attributes in a table
            if len(tbl.relationships) > 0:
                document.add_heading('Relaciones', level=3)

                rel_table = document.add_table(rows=len(tbl.relationships) + 1, cols=3)

                hdr_cells = rel_table.rows[0].cells
                hdr_cells[0].text = 'Atributo(s)'
                hdr_cells[1].text = 'Tabla relacionada'
                hdr_cells[2].text = 'Atributo(s) de la tabla relacionada'

                for i, rel in enumerate(tbl.relationships):
                    af = ','.join([e.name for e in rel.attributes_from])
                    at = ','.join([e.name for e in rel.attributes_to])

                    row_cells = rel_table.rows[i + 1].cells
                    row_cells[0].text = af
                    row_cells[1].text = rel.tbl_to.name
                    row_cells[2].text = at

            # add a page break
            document.add_page_break()

        if file_name is None:
            file_name = self.output_file.replace('.sql', '.docx')
        document.save(file_name)

    def to_excel(self, file_name=None):
        """
        Write this model to excel
        """
        if file_name is None:
            file_name = self.output_file.replace('.sql', '.xlsx')

        writer = pd.ExcelWriter(file_name)

        for tbl in self.tables:

            attr_names = [a.name for a in tbl.attributes]
            data = [''] * len(attr_names)
            df = pd.DataFrame(data=[data], columns=attr_names)
            df.to_excel(writer, sheet_name=tbl.name[:30], index=False)

        writer.save()
        writer.close()
    
    def save(self, file_name='sql.txt'):
        text_file = open(self.output_file, "w")
        text_file.write(self.sql_code)
        text_file.close()


def model_to_ms_word(model, file_name=None, table_style='Plain Table 1'):
    """
    Export the Model to a MS Word document for Trazar, else use the to_ms_word method of the model
    :param: file_name: Name of the file
    """
    document = Document('plantilla_sql2dia.docx')

    # group tables by the comment
    categories = {'otras': list()}
    for tbl in model.tables:
        c = tbl.comment.split(':')
        if len(c) < 2:
            categories['otras'].append(tbl)
        else:
            cat = c[0]
            if cat not in categories.keys():
                categories[cat] = list()
            categories[cat].append(tbl)

    # sort the categories
    dcategories = list(categories.keys())
    dcategories.sort()

    for cat in dcategories:
        tables = categories[cat]

        if len(tables) > 0:
            document.add_heading(cat, level=1)

            for tbl in tables:

                # one section per table
                document.add_heading(tbl.name, level=2)

                # add the comment
                p = document.add_paragraph('\n' + tbl.comment + '\n')

                document.add_heading('Atributos', level=3)

                # add the attributes in a table
                doc_table = document.add_table(rows=len(tbl.attributes) + 1, cols=5, style=table_style)

                hdr_cells = doc_table.rows[0].cells
                hdr_cells[0].text = 'Atributo'
                hdr_cells[1].text = 'Tipo'
                hdr_cells[2].text = 'Clave Primaria'
                hdr_cells[3].text = 'Único'
                hdr_cells[4].text = 'Comentario'

                for i, attr in enumerate(tbl.attributes):
                    row_cells = doc_table.rows[i+1].cells
                    row_cells[0].text = attr.name  # 'Atributo'
                    row_cells[1].text = str(attr.type)  # 'Tipo'
                    row_cells[2].text = str(attr.is_primary_key)  # 'Clave Primaria'
                    row_cells[3].text = str(attr.is_unique)  # 'Único'
                    row_cells[4].text = attr.comment  # 'Comentario'

                # add the attributes in a table
                if len(tbl.relationships) > 0:
                    document.add_heading('Relaciones', level=3)

                    rel_table = document.add_table(rows=len(tbl.relationships) + 1, cols=3, style=table_style)

                    hdr_cells = rel_table.rows[0].cells
                    hdr_cells[0].text = 'Atributo(s)'
                    hdr_cells[1].text = 'Tabla relacionada'
                    hdr_cells[2].text = 'Atributo(s) de la tabla relacionada'

                    for i, rel in enumerate(tbl.relationships):
                        af = ','.join([e.name for e in rel.attributes_from])
                        at = ','.join([e.name for e in rel.attributes_to])

                        row_cells = rel_table.rows[i + 1].cells
                        row_cells[0].text = af
                        row_cells[1].text = rel.tbl_to.name
                        row_cells[2].text = at

                # add a page break
                document.add_page_break()

    if file_name is None:
        file_name = model.output_file.replace('.sql', '.docx')
    document.save(file_name)


if __name__ == '__main__':

    fname = r'trazar v10.dia'
    # fname = 'Diagrama_PMV3.dia'
    model = DBModel(fname)
    print(model.to_sql())
    model.save()

    # model.to_ms_word()
    # model.to_excel()
    # model_to_ms_word(model)
